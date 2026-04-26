import os
import fitz
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,
    chunk_overlap=100
)

def load_pdf(file_path: str):
    doc = fitz.open(file_path)
    documents = []
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            documents.append(Document(
                page_content=text,
                metadata={
                    "source_file": os.path.basename(file_path),
                    "page_number": page_num + 1
                }
            ))
    return documents

def load_docx(file_path: str):
    doc = DocxDocument(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(
        page_content=full_text,
        metadata={"source_file": os.path.basename(file_path), "page_number": 1}
    )]

def ingest_folder(folder_path: str):
    all_chunks = []
    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            raw_docs = load_pdf(filepath)
        elif filename.endswith(".docx"):
            raw_docs = load_docx(filepath)
        else:
            continue
        chunks = splitter.split_documents(raw_docs)
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i
        all_chunks.extend(chunks)
        print(f"Loaded {len(chunks)} chunks from {filename}")
    return all_chunks